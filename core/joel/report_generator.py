"""
Joel Report Generator v2 — Relatórios Profissionais Nível RI

Gera relatórios com qualidade de Relação com Investidores:
- PDF: Capa profissional, sumário, header/footer, gráficos embutidos, tabelas estilizadas
- DOCX: Estilos corporativos, gráficos, tabelas formatadas, header/footer
- XLSX: Dados estruturados com formatação condicional
- TXT: Plain text limpo com boa formatação
- HTML: Markdown + CSS profissional com gráficos inline

Inspirado em relatórios de RI de empresas Blue Chip (Itaú, Vale, Petrobras, etc.)
"""

import os
import io
import re
import base64
import logging
from datetime import datetime
from typing import Optional

import markdown as md
import bleach
from django.conf import settings

logger = logging.getLogger(__name__)

# ============================================================
# CORPORATE THEME
# ============================================================
THEME = {
    "primary": "#1e3a5f",
    "secondary": "#2563eb",
    "accent": "#7c3aed",
    "success": "#059669",
    "warning": "#d97706",
    "danger": "#dc2626",
    "text": "#1f2937",
    "text_light": "#6b7280",
    "border": "#e5e7eb",
    "bg": "#ffffff",
    "bg_alt": "#f8fafc",
    "bg_section": "#f0f4ff",
}

ALLOWED_TAGS = bleach.ALLOWED_TAGS | {
    "h1", "h2", "h3", "h4", "h5", "h6", "p", "br", "hr",
    "table", "thead", "tbody", "tr", "th", "td",
    "ul", "ol", "li", "strong", "em", "a", "code", "pre", "blockquote",
    "img", "div", "span", "figure", "figcaption",
}
ALLOWED_ATTRS = {
    "a": ["href", "title", "target"],
    "img": ["src", "alt", "width", "style", "class"],
    "td": ["colspan", "rowspan", "style"],
    "th": ["colspan", "rowspan", "style"],
    "div": ["class", "style"],
    "span": ["class", "style"],
    "table": ["class", "style"],
}


def markdown_to_html(markdown_text: str, charts_base64: list[dict] = None) -> str:
    """
    Converte Markdown para HTML sanitizado com estilo profissional.
    Opcionalmente embute gráficos base64 no HTML.
    """
    html = md.markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "toc", "nl2br", "sane_lists"],
    )
    
    # Wrap tables in responsive container
    html = html.replace("<table>", '<div class="table-wrapper"><table class="report-table">')
    html = html.replace("</table>", "</table></div>")
    
    # Add chart images if provided
    if charts_base64:
        chart_html = '<div class="charts-section">'
        chart_html += '<h2 class="charts-header">Visualizações</h2>'
        for chart in charts_base64:
            chart_html += f'''
            <figure class="chart-figure">
                <img src="data:image/png;base64,{chart['base64']}" 
                     alt="{chart.get('title', 'Gráfico')}"
                     class="chart-image" />
                <figcaption>{chart.get('title', '')}</figcaption>
            </figure>'''
        chart_html += '</div>'
        
        # Insert charts before "Conclus" section or at end
        conclus_patterns = [
            r'<h[23][^>]*>.*?[Cc]onclus',
            r'<h[23][^>]*>.*?[Rr]ecomenda',
            r'<h[23][^>]*>.*?[Cc]onsider',
        ]
        inserted = False
        for pattern in conclus_patterns:
            match = re.search(pattern, html)
            if match:
                html = html[:match.start()] + chart_html + html[match.start():]
                inserted = True
                break
        
        if not inserted:
            html += chart_html
    
    return bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS)


def generate_pdf(
    markdown_text: str,
    title: str = "Relatório",
    charts_base64: list[dict] = None,
    professional_area: str = "",
    report_type: str = "",
) -> io.BytesIO:
    """
    Gera PDF profissional nível RI (Relação com Investidores).
    
    Features:
    - Capa profissional com gradient
    - Header/footer em todas as páginas
    - Sumário automático
    - Gráficos embutidos
    - Tabelas estilizadas
    - Tipografia profissional
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor, Color
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        Table, TableStyle, HRFlowable, Image, KeepTogether,
    )
    from reportlab.platypus.flowables import Flowable
    
    buffer = io.BytesIO()
    
    # Page dimensions
    page_w, page_h = A4
    margin_lr = 2.2 * cm
    margin_top = 2.8 * cm
    margin_bottom = 2.2 * cm
    
    # === Custom flowables ===
    class GradientBar(Flowable):
        """Full-width gradient bar."""
        def __init__(self, width, height=4, colors=None):
            Flowable.__init__(self)
            self.width = width
            self.height = height
            self.colors = colors or [HexColor(THEME["primary"]), HexColor(THEME["secondary"])]
        
        def draw(self):
            steps = 100
            step_w = self.width / steps
            c1, c2 = self.colors
            for i in range(steps):
                r = c1.red + (c2.red - c1.red) * i / steps
                g = c1.green + (c2.green - c1.green) * i / steps
                b = c1.blue + (c2.blue - c1.blue) * i / steps
                self.canv.setFillColor(Color(r, g, b))
                self.canv.rect(i * step_w, 0, step_w + 1, self.height, fill=1, stroke=0)
    
    class CoverPage(Flowable):
        """Professional cover page."""
        def __init__(self, title, area, rtype, date_str):
            Flowable.__init__(self)
            self.title = title
            self.area = area
            self.rtype = rtype
            self.date_str = date_str
            self.width = page_w - 2 * margin_lr
            self.height = page_h - margin_top - margin_bottom
        
        def draw(self):
            c = self.canv
            w, h = self.width, self.height
            
            # Top gradient bar
            steps = 150
            bar_h = 8
            step_w = w / steps
            c1 = HexColor(THEME["primary"])
            c2 = HexColor(THEME["secondary"])
            for i in range(steps):
                r = c1.red + (c2.red - c1.red) * i / steps
                g = c1.green + (c2.green - c1.green) * i / steps
                b = c1.blue + (c2.blue - c1.blue) * i / steps
                c.setFillColor(Color(r, g, b))
                c.rect(i * step_w, h - 20, step_w + 1, bar_h, fill=1, stroke=0)
            
            # Title block
            c.setFillColor(HexColor(THEME["primary"]))
            c.setFont("Helvetica-Bold", 26)
            
            # Word wrap title
            title_lines = self._wrap_text(self.title, 26, w - 20)
            y_pos = h - 100
            for tl in title_lines:
                c.drawString(10, y_pos, tl)
                y_pos -= 36
            
            # Accent line under title
            y_pos -= 10
            c.setStrokeColor(HexColor(THEME["secondary"]))
            c.setLineWidth(3)
            c.line(10, y_pos, 180, y_pos)
            
            # Area & Type
            y_pos -= 40
            c.setFillColor(HexColor(THEME["secondary"]))
            c.setFont("Helvetica", 13)
            if self.area:
                c.drawString(10, y_pos, f"Área: {self.area}")
                y_pos -= 22
            if self.rtype:
                c.drawString(10, y_pos, f"Tipo: {self.rtype}")
                y_pos -= 22
            
            # Date
            y_pos -= 20
            c.setFillColor(HexColor(THEME["text_light"]))
            c.setFont("Helvetica", 11)
            c.drawString(10, y_pos, self.date_str)
            
            # Bottom corner branding
            c.setFillColor(HexColor(THEME["primary"]))
            c.setFont("Helvetica-Bold", 14)
            c.drawString(10, 40, "Joel")
            c.setFillColor(HexColor(THEME["text_light"]))
            c.setFont("Helvetica", 9)
            c.drawString(10, 25, "Agente de Análise de Documentos")
            
            # Bottom gradient bar
            for i in range(steps):
                r = c2.red + (c1.red - c2.red) * i / steps
                g = c2.green + (c1.green - c2.green) * i / steps
                b = c2.blue + (c1.blue - c2.blue) * i / steps
                c.setFillColor(Color(r, g, b))
                c.rect(i * step_w, 0, step_w + 1, bar_h, fill=1, stroke=0)
        
        def _wrap_text(self, text, font_size, max_width):
            """Simple word wrap."""
            from reportlab.pdfbase.pdfmetrics import stringWidth
            words = text.split()
            lines, current = [], ""
            for w in words:
                test = f"{current} {w}".strip()
                if stringWidth(test, "Helvetica-Bold", font_size) <= max_width:
                    current = test
                else:
                    if current:
                        lines.append(current)
                    current = w
            if current:
                lines.append(current)
            return lines[:3]  # Max 3 lines
    
    # === Header & Footer ===
    page_counter = [0]
    
    def header_footer(canvas, doc):
        page_counter[0] += 1
        page_num = page_counter[0]
        
        if page_num <= 1:
            return  # No header/footer on cover
        
        canvas.saveState()
        
        # Header line
        canvas.setStrokeColor(HexColor(THEME["border"]))
        canvas.setLineWidth(0.5)
        canvas.line(margin_lr, page_h - margin_top + 10, page_w - margin_lr, page_h - margin_top + 10)
        
        # Header text
        canvas.setFillColor(HexColor(THEME["text_light"]))
        canvas.setFont("Helvetica", 7)
        canvas.drawString(margin_lr, page_h - margin_top + 15, "Joel — Análise Profissional de Documentos")
        canvas.drawRightString(page_w - margin_lr, page_h - margin_top + 15, title[:50])
        
        # Footer
        canvas.setStrokeColor(HexColor(THEME["border"]))
        canvas.line(margin_lr, margin_bottom - 10, page_w - margin_lr, margin_bottom - 10)
        
        canvas.setFillColor(HexColor(THEME["text_light"]))
        canvas.setFont("Helvetica", 7)
        canvas.drawString(margin_lr, margin_bottom - 22, 
                         f"Confidencial — Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        canvas.drawRightString(page_w - margin_lr, margin_bottom - 22, f"Página {page_num}")
        
        # Bottom accent bar
        canvas.setFillColor(HexColor(THEME["secondary"]))
        canvas.rect(margin_lr, margin_bottom - 28, 30, 2, fill=1, stroke=0)
        
        canvas.restoreState()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=margin_lr,
        leftMargin=margin_lr,
        topMargin=margin_top,
        bottomMargin=margin_bottom,
    )
    
    styles = getSampleStyleSheet()
    content_width = page_w - 2 * margin_lr
    
    # === Custom Styles ===
    styles.add(ParagraphStyle(
        name="CoverTitle", parent=styles["Title"],
        fontSize=26, textColor=HexColor(THEME["primary"]),
        spaceAfter=10, alignment=TA_LEFT, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="TOCTitle", parent=styles["Heading1"],
        fontSize=18, textColor=HexColor(THEME["primary"]),
        spaceBefore=0, spaceAfter=20, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="TOCEntry", parent=styles["Normal"],
        fontSize=10, textColor=HexColor(THEME["text"]),
        leftIndent=15, spaceBefore=4, spaceAfter=4,
        fontName="Helvetica",
    ))
    styles.add(ParagraphStyle(
        name="SectionH1", parent=styles["Heading1"],
        fontSize=18, textColor=HexColor(THEME["primary"]),
        spaceBefore=24, spaceAfter=8, fontName="Helvetica-Bold",
        borderWidth=0, borderPadding=0,
    ))
    styles.add(ParagraphStyle(
        name="SectionH2", parent=styles["Heading2"],
        fontSize=14, textColor=HexColor(THEME["secondary"]),
        spaceBefore=18, spaceAfter=6, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="SectionH3", parent=styles["Heading3"],
        fontSize=12, textColor=HexColor(THEME["accent"]),
        spaceBefore=12, spaceAfter=4, fontName="Helvetica-Bold",
    ))
    styles.add(ParagraphStyle(
        name="BodyText2", parent=styles["Normal"],
        fontSize=10, leading=15, alignment=TA_JUSTIFY,
        spaceAfter=6, fontName="Helvetica",
        textColor=HexColor(THEME["text"]),
    ))
    styles.add(ParagraphStyle(
        name="BulletText", parent=styles["Normal"],
        fontSize=10, leading=14, leftIndent=20,
        spaceAfter=3, fontName="Helvetica",
        textColor=HexColor(THEME["text"]),
        bulletIndent=8,
    ))
    styles.add(ParagraphStyle(
        name="NumberText", parent=styles["Normal"],
        fontSize=10, leading=14, leftIndent=20,
        spaceAfter=3, fontName="Helvetica",
        textColor=HexColor(THEME["text"]),
        bulletIndent=8,
    ))
    styles.add(ParagraphStyle(
        name="BlockQuote", parent=styles["Normal"],
        fontSize=10, leading=14, leftIndent=25, rightIndent=15,
        spaceAfter=8, spaceBefore=8, fontName="Helvetica-Oblique",
        textColor=HexColor(THEME["text_light"]),
        borderWidth=2, borderColor=HexColor(THEME["secondary"]),
        borderPadding=8,
    ))
    styles.add(ParagraphStyle(
        name="FooterNote", parent=styles["Normal"],
        fontSize=8, textColor=HexColor(THEME["text_light"]),
        alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        name="ChartCaption", parent=styles["Normal"],
        fontSize=9, textColor=HexColor(THEME["text_light"]),
        alignment=TA_CENTER, spaceAfter=12, spaceBefore=4,
        fontName="Helvetica-Oblique",
    ))
    
    story = []
    
    # === 1. COVER PAGE ===
    date_str = datetime.now().strftime("%d de %B de %Y").replace(
        "January", "Janeiro").replace("February", "Fevereiro").replace(
        "March", "Março").replace("April", "Abril").replace(
        "May", "Maio").replace("June", "Junho").replace(
        "July", "Julho").replace("August", "Agosto").replace(
        "September", "Setembro").replace("October", "Outubro").replace(
        "November", "Novembro").replace("December", "Dezembro")
    
    cover = CoverPage(title, professional_area, report_type, date_str)
    story.append(cover)
    story.append(PageBreak())
    
    # === 2. TABLE OF CONTENTS ===
    lines = markdown_text.split("\n")
    toc_items = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## ") and not stripped.startswith("### "):
            toc_items.append({"level": 2, "text": stripped[3:].strip()})
        elif stripped.startswith("### "):
            toc_items.append({"level": 3, "text": stripped[4:].strip()})
    
    if toc_items:
        story.append(Paragraph("Sumário", styles["TOCTitle"]))
        story.append(GradientBar(content_width, 3))
        story.append(Spacer(1, 0.5 * cm))
        
        for idx, item in enumerate(toc_items, 1):
            prefix = "    " if item["level"] == 3 else ""
            clean_text = re.sub(r'\*\*|__', '', item["text"])
            
            if item["level"] == 2:
                entry_style = ParagraphStyle(
                    f"toc_entry_{idx}", parent=styles["TOCEntry"],
                    fontName="Helvetica-Bold", fontSize=10,
                    textColor=HexColor(THEME["primary"]),
                )
            else:
                entry_style = ParagraphStyle(
                    f"toc_sub_{idx}", parent=styles["TOCEntry"],
                    leftIndent=35, fontSize=9,
                    textColor=HexColor(THEME["text_light"]),
                )
            
            story.append(Paragraph(f"{prefix}{clean_text}", entry_style))
        
        story.append(Spacer(1, 0.5 * cm))
        story.append(GradientBar(content_width, 2))
        story.append(PageBreak())
    
    # === 3. MAIN CONTENT ===
    in_table = False
    table_rows = []
    in_blockquote = False
    blockquote_lines = []
    number_counter = 0
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Empty line
        if not stripped:
            if in_blockquote and blockquote_lines:
                bq_text = "<br/>".join(blockquote_lines)
                story.append(Paragraph(bq_text, styles["BlockQuote"]))
                blockquote_lines = []
                in_blockquote = False
            story.append(Spacer(1, 0.2 * cm))
            number_counter = 0
            i += 1
            continue
        
        # Blockquote
        if stripped.startswith(">"):
            in_blockquote = True
            blockquote_lines.append(stripped.lstrip("> ").strip())
            i += 1
            continue
        elif in_blockquote and blockquote_lines:
            bq_text = "<br/>".join(blockquote_lines)
            story.append(Paragraph(bq_text, styles["BlockQuote"]))
            blockquote_lines = []
            in_blockquote = False
        
        # Table detection
        if stripped.startswith("|") and "|" in stripped[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            
            # Parse table row
            cells = [c.strip() for c in stripped.split("|")]
            cells = [c for c in cells if c]
            
            # Skip separator rows
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            
            table_rows.append(cells)
            i += 1
            continue
        elif in_table and table_rows:
            story.append(_build_pdf_table(table_rows, styles, content_width))
            story.append(Spacer(1, 0.3 * cm))
            table_rows = []
            in_table = False
        
        # Headings
        if stripped.startswith("#### "):
            text = _clean_md(stripped[5:])
            h4_style = ParagraphStyle(
                "h4_inline", parent=styles["SectionH3"],
                fontSize=11, textColor=HexColor(THEME["text"]),
                spaceBefore=10,
            )
            story.append(Paragraph(text, h4_style))
            i += 1
            continue
        elif stripped.startswith("### "):
            text = _clean_md(stripped[4:])
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph(text, styles["SectionH3"]))
            i += 1
            continue
        elif stripped.startswith("## "):
            text = _clean_md(stripped[3:])
            story.append(Spacer(1, 0.4 * cm))
            story.append(GradientBar(content_width, 2))
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(text, styles["SectionH2"]))
            i += 1
            continue
        elif stripped.startswith("# "):
            text = _clean_md(stripped[2:])
            story.append(Spacer(1, 0.5 * cm))
            story.append(GradientBar(content_width, 3))
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph(text, styles["SectionH1"]))
            i += 1
            continue
        
        # Horizontal rule
        if stripped == "---" or stripped == "***" or stripped == "___":
            story.append(Spacer(1, 0.2 * cm))
            story.append(GradientBar(content_width, 1.5,
                                     [HexColor(THEME["border"]), HexColor("#ffffff")]))
            story.append(Spacer(1, 0.2 * cm))
            i += 1
            continue
        
        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = _format_inline(stripped[2:])
            story.append(Paragraph(f"• {text}", styles["BulletText"]))
            i += 1
            continue
        
        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            number_counter += 1
            text = _format_inline(num_match.group(2))
            story.append(Paragraph(f"{number_counter}. {text}", styles["NumberText"]))
            i += 1
            continue
        
        # Regular paragraph
        text = _format_inline(stripped)
        try:
            story.append(Paragraph(text, styles["BodyText2"]))
        except Exception:
            clean = bleach.clean(stripped, tags=set(), strip=True)
            story.append(Paragraph(clean, styles["BodyText2"]))
        
        i += 1
    
    # Flush remaining table
    if in_table and table_rows:
        story.append(_build_pdf_table(table_rows, styles, content_width))
    if in_blockquote and blockquote_lines:
        bq_text = "<br/>".join(blockquote_lines)
        story.append(Paragraph(bq_text, styles["BlockQuote"]))
    
    # === 4. EMBEDDED CHARTS ===
    if charts_base64:
        story.append(Spacer(1, 0.5 * cm))
        story.append(GradientBar(content_width, 3))
        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph("Visualizações", styles["SectionH2"]))
        story.append(Spacer(1, 0.3 * cm))
        
        for chart in charts_base64:
            try:
                img_data = base64.b64decode(chart["base64"])
                img_buf = io.BytesIO(img_data)
                img = Image(img_buf, width=content_width * 0.85, height=content_width * 0.45)
                story.append(KeepTogether([
                    img,
                    Paragraph(chart.get("title", ""), styles["ChartCaption"]),
                    Spacer(1, 0.4 * cm),
                ]))
            except Exception as e:
                logger.warning(f"Erro ao embutir gráfico no PDF: {e}")
    
    # === 5. FOOTER DISCLAIMER ===
    story.append(Spacer(1, 1 * cm))
    story.append(GradientBar(content_width, 2))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        "<i>Este relatório foi gerado automaticamente por Joel — Agente de Análise de Documentos. "
        "As informações e análises contidas neste documento são baseadas nos dados fornecidos e em "
        "fontes públicas disponíveis. Recomenda-se validação independente antes de decisões críticas.</i>",
        styles["FooterNote"],
    ))
    
    try:
        doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)
    except Exception as e:
        logger.error(f"Erro ao gerar PDF: {e}", exc_info=True)
        # Fallback simples
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        fallback_styles = getSampleStyleSheet()
        fallback_story = [
            Paragraph(title, fallback_styles["Title"]),
            Spacer(1, 1 * cm),
        ]
        for line in markdown_text.split("\n"):
            if line.strip():
                try:
                    fallback_story.append(Paragraph(
                        bleach.clean(line.strip(), tags=set(), strip=True),
                        fallback_styles["Normal"]
                    ))
                except Exception:
                    pass
        doc.build(fallback_story)
    
    buffer.seek(0)
    return buffer


def _build_pdf_table(rows: list[list[str]], styles, max_width: float):
    """Build a professionally styled ReportLab table from parsed rows."""
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import cm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
    
    if not rows:
        return Spacer(1, 0)
    
    # Determine column count
    max_cols = max(len(r) for r in rows)
    
    # Normalize rows
    normalized = []
    for row in rows:
        while len(row) < max_cols:
            row.append("")
        normalized.append(row)
    
    # Wrap cell content in Paragraphs
    cell_style = ParagraphStyle(
        "TableCell", parent=styles["Normal"],
        fontSize=9, leading=11,
        textColor=HexColor(THEME["text"]),
    )
    header_cell_style = ParagraphStyle(
        "TableHeaderCell", parent=cell_style,
        fontName="Helvetica-Bold",
        textColor=HexColor("#ffffff"),
        fontSize=9,
    )
    
    table_data = []
    for ridx, row in enumerate(normalized):
        table_row = []
        for cell in row:
            style = header_cell_style if ridx == 0 else cell_style
            clean = _clean_md(cell)
            table_row.append(Paragraph(clean, style))
        table_data.append(table_row)
    
    # Calculate column widths
    col_width = min(max_width / max_cols, 5 * cm)
    col_widths = [col_width] * max_cols
    if max_cols > 1:
        total = max_width * 0.95
        col_widths[0] = total * 0.35
        remaining = total - col_widths[0]
        for c in range(1, max_cols):
            col_widths[c] = remaining / (max_cols - 1)
    
    table = Table(table_data, colWidths=col_widths)
    
    # Professional table styling
    table_style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(THEME["primary"])),
        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("TOPPADDING", (0, 0), (-1, 0), 8),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 1), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 6),
        ("TOPPADDING", (0, 1), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor(THEME["border"])),
    ])
    
    # Alternate row colors
    for ridx in range(1, len(table_data)):
        bg = HexColor(THEME["bg_alt"]) if ridx % 2 == 0 else HexColor(THEME["bg"])
        table_style.add("BACKGROUND", (0, ridx), (-1, ridx), bg)
    
    table.setStyle(table_style)
    return table


def _format_inline(text: str) -> str:
    """Convert markdown inline formatting to ReportLab XML tags."""
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.+?)__', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    text = re.sub(r'_(.+?)_', r'<i>\1</i>', text)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2" color="#2563eb">\1</a>', text)
    text = re.sub(r'`(.+?)`', r'<font face="Courier" size="9" color="#7c3aed">\1</font>', text)
    return text


def _clean_md(text: str) -> str:
    """Remove markdown formatting characters."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def generate_docx(
    markdown_text: str,
    title: str = "Relatório",
    charts_base64: list[dict] = None,
    professional_area: str = "",
    report_type: str = "",
) -> io.BytesIO:
    """
    Gera DOCX profissional com estilos corporativos.
    
    Features:
    - Custom heading styles (brand colors)
    - Professional tables
    - Embedded charts
    - Header/footer with page numbers
    - Cover page
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    
    doc = Document()
    
    # === Configure styles ===
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.line_spacing = 1.15
    
    # Heading styles
    for level, (size, color_hex) in enumerate([
        (22, THEME["primary"]),
        (16, THEME["secondary"]),
        (13, THEME["accent"]),
    ], 1):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            hs = doc.styles[style_name]
            hs.font.size = Pt(size)
            r, g, b = int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16)
            hs.font.color.rgb = RGBColor(r, g, b)
            hs.font.bold = True
            hs.paragraph_format.space_before = Pt(18 if level == 1 else 14)
            hs.paragraph_format.space_after = Pt(8)
    
    # === COVER PAGE ===
    for _ in range(4):
        doc.add_paragraph("")
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    run.bold = True
    
    # Accent line
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = accent_table.cell(0, 0)
    cell.text = ""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2563eb"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    tr = accent_table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="60" w:hRule="exact"/>')
    trPr.append(trHeight)
    
    doc.add_paragraph("")
    
    # Metadata
    meta_items = []
    if professional_area:
        meta_items.append(f"Área: {professional_area}")
    if report_type:
        meta_items.append(f"Tipo: {report_type}")
    meta_items.append(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    for item in meta_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Joel")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("Agente de Análise de Documentos")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    doc.add_page_break()
    
    # === HEADER / FOOTER ===
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"Joel — {title[:40]}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Confidencial — Joel Análise de Documentos — Página ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    # Page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run2 = fp.add_run()
    run2._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run3 = fp.add_run()
    run3._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run4 = fp.add_run()
    run4._r.append(fldChar2)
    
    # === TABLE OF CONTENTS ===
    lines = markdown_text.split("\n")
    toc_items = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## ") and not stripped.startswith("### "):
            toc_items.append({"level": 2, "text": _clean_md(stripped[3:])})
        elif stripped.startswith("### "):
            toc_items.append({"level": 3, "text": _clean_md(stripped[4:])})
    
    if toc_items:
        doc.add_heading("Sumário", level=1)
        for item in toc_items:
            p = doc.add_paragraph()
            prefix = "   " if item["level"] == 3 else ""
            bullet = "›" if item["level"] == 3 else "■"
            run = p.add_run(f"{prefix}{bullet}  {item['text']}")
            run.font.size = Pt(10 if item["level"] == 2 else 9)
            if item["level"] == 2:
                run.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            else:
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        doc.add_page_break()
    
    # === MAIN CONTENT ===
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        if not stripped:
            i += 1
            continue
        
        # Table detection
        if stripped.startswith("|") and "|" in stripped[1:]:
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in stripped.split("|")]
            cells = [c for c in cells if c]
            if not all(re.match(r'^[-:]+$', c) for c in cells):
                table_rows.append(cells)
            i += 1
            continue
        elif in_table and table_rows:
            _add_docx_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        # Headings
        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(_clean_md(stripped[5:]))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        elif stripped.startswith("### "):
            doc.add_heading(_clean_md(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(_clean_md(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(_clean_md(stripped[2:]), level=1)
        elif stripped == "---" or stripped == "***":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)
            run.font.size = Pt(6)
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(_clean_md(stripped[2:]))
            run.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(_clean_md(stripped[2:]), style="List Bullet")
        elif re.match(r'^\d+\.\s+', stripped):
            num_match = re.match(r'^\d+\.\s+(.+)', stripped)
            if num_match:
                doc.add_paragraph(_clean_md(num_match.group(1)), style="List Number")
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)
        
        i += 1
    
    # Flush remaining table
    if in_table and table_rows:
        _add_docx_table(doc, table_rows)
    
    # === EMBEDDED CHARTS ===
    if charts_base64:
        doc.add_page_break()
        doc.add_heading("Visualizações", level=1)
        
        for chart in charts_base64:
            try:
                img_data = base64.b64decode(chart["base64"])
                img_stream = io.BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(5.5))
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(chart.get("title", ""))
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            except Exception as e:
                logger.warning(f"Erro ao embutir gráfico no DOCX: {e}")
    
    # === DISCLAIMER ===
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "Este relatório foi gerado automaticamente por Joel — Agente de Análise de Documentos. "
        "As informações e análises contidas neste documento são baseadas nos dados fornecidos e em "
        "fontes públicas disponíveis. Recomenda-se validação independente antes de decisões críticas."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_docx_table(doc, rows: list[list[str]]):
    """Add a professionally styled table to the DOCX document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    if not rows:
        return
    
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    
    for ridx, row in enumerate(rows):
        for cidx, cell_text in enumerate(row):
            if cidx < max_cols:
                cell = table.cell(ridx, cidx)
                cell.text = _clean_md(cell_text)
                
                if ridx == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1e3a5f"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                            run.font.size = Pt(9)
                elif ridx % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f8fafc"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


def _add_formatted_runs(paragraph, text: str):
    """Add formatted runs to a paragraph, handling bold/italic/links."""
    from docx.shared import Pt, RGBColor
    
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|__(.+?)__|\*(.+?)\*|_(.+?)_|\[([^\]]+)\]\(([^)]+)\)|`(.+?)`|([^*_\[`]+))'
    
    for match in re.finditer(pattern, text):
        groups = match.groups()
        run = None
        
        if groups[1]:  # ***bold italic***
            run = paragraph.add_run(groups[1])
            run.bold = True
            run.italic = True
        elif groups[2]:  # **bold**
            run = paragraph.add_run(groups[2])
            run.bold = True
        elif groups[3]:  # __bold__
            run = paragraph.add_run(groups[3])
            run.bold = True
        elif groups[4]:  # *italic*
            run = paragraph.add_run(groups[4])
            run.italic = True
        elif groups[5]:  # _italic_
            run = paragraph.add_run(groups[5])
            run.italic = True
        elif groups[6]:  # [text](url)
            run = paragraph.add_run(groups[6])
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.underline = True
        elif groups[8]:  # `code`
            run = paragraph.add_run(groups[8])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        elif groups[9]:  # plain text
            run = paragraph.add_run(groups[9])
        
        if run:
            if not run.font.size:
                run.font.size = Pt(11)


def generate_xlsx(
    markdown_text: str,
    references: list = None,
    title: str = "Relatório",
    charts_base64: list[dict] = None,
) -> io.BytesIO:
    """
    Gera XLSX com dados estruturados, formatação condicional e gráficos.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.drawing.image import Image as XlImage
    
    wb = Workbook()
    
    thin_border = Border(
        left=Side(style="thin", color="E5E7EB"),
        right=Side(style="thin", color="E5E7EB"),
        top=Side(style="thin", color="E5E7EB"),
        bottom=Side(style="thin", color="E5E7EB"),
    )
    header_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=10, name="Calibri")
    alt_fill = PatternFill(start_color="F8FAFC", end_color="F8FAFC", fill_type="solid")
    
    # === Sheet 1: Relatório ===
    ws = wb.active
    ws.title = "Relatório"
    ws.sheet_properties.tabColor = "1E3A5F"
    
    ws.merge_cells("A1:D1")
    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=18, color="1E3A5F", name="Calibri")
    ws["A1"].alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 40
    
    ws["A2"] = f"Gerado por Joel — {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws["A2"].font = Font(italic=True, size=9, color="6B7280", name="Calibri")
    ws.row_dimensions[2].height = 20
    
    for col in range(1, 5):
        cell = ws.cell(row=3, column=col)
        cell.fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
    ws.row_dimensions[3].height = 4
    
    row = 5
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            row += 1
            continue
        
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            continue
        
        cell = ws.cell(row=row, column=1)
        clean = _clean_md(stripped.lstrip("#").strip())
        cell.value = clean
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        
        if stripped.startswith("# ") and not stripped.startswith("## "):
            cell.font = Font(bold=True, size=16, color="1E3A5F", name="Calibri")
            ws.row_dimensions[row].height = 30
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            cell.font = Font(bold=True, size=13, color="2563EB", name="Calibri")
            ws.row_dimensions[row].height = 26
            for col in range(1, 5):
                ws.cell(row=row, column=col).fill = PatternFill(
                    start_color="F0F4FF", end_color="F0F4FF", fill_type="solid"
                )
        elif stripped.startswith("### "):
            cell.font = Font(bold=True, size=11, color="7C3AED", name="Calibri")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            cell.value = f"  •  {clean}"
            cell.font = Font(size=10, name="Calibri")
        elif stripped.startswith("|"):
            cells_data = [c.strip() for c in stripped.split("|") if c.strip()]
            for cidx, cval in enumerate(cells_data):
                col_cell = ws.cell(row=row, column=cidx + 1, value=_clean_md(cval))
                col_cell.font = Font(size=9, name="Calibri")
                col_cell.border = thin_border
                col_cell.alignment = Alignment(wrap_text=True)
        else:
            cell.font = Font(size=10, name="Calibri")
        
        row += 1
    
    ws.column_dimensions["A"].width = 80
    ws.column_dimensions["B"].width = 25
    ws.column_dimensions["C"].width = 25
    ws.column_dimensions["D"].width = 25
    
    # === Sheet 2: Referências ===
    if references:
        ws_ref = wb.create_sheet("Referências")
        ws_ref.sheet_properties.tabColor = "2563EB"
        
        headers = ["#", "Título", "URL", "Resumo"]
        for col, header_text in enumerate(headers, 1):
            cell = ws_ref.cell(row=1, column=col, value=header_text)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center")
        
        for idx, ref in enumerate(references, 1):
            row_num = idx + 1
            ws_ref.cell(row=row_num, column=1, value=idx).border = thin_border
            ws_ref.cell(row=row_num, column=2, value=ref.get("title", "")).border = thin_border
            
            url_cell = ws_ref.cell(row=row_num, column=3, value=ref.get("url", ""))
            url_cell.font = Font(color="2563EB", underline="single", size=9, name="Calibri")
            url_cell.border = thin_border
            
            content_cell = ws_ref.cell(row=row_num, column=4, value=ref.get("content", "")[:500])
            content_cell.alignment = Alignment(wrap_text=True)
            content_cell.border = thin_border
            
            if idx % 2 == 0:
                for col in range(1, 5):
                    ws_ref.cell(row=row_num, column=col).fill = alt_fill
        
        ws_ref.column_dimensions["A"].width = 5
        ws_ref.column_dimensions["B"].width = 40
        ws_ref.column_dimensions["C"].width = 55
        ws_ref.column_dimensions["D"].width = 70
        ws_ref.auto_filter.ref = f"A1:D{len(references) + 1}"
    
    # === Sheet 3: Charts ===
    if charts_base64:
        ws_charts = wb.create_sheet("Gráficos")
        ws_charts.sheet_properties.tabColor = "7C3AED"
        
        ws_charts["A1"] = "Visualizações do Relatório"
        ws_charts["A1"].font = Font(bold=True, size=14, color="1E3A5F", name="Calibri")
        
        chart_row = 3
        for chart in charts_base64:
            try:
                img_data = base64.b64decode(chart["base64"])
                img_buf = io.BytesIO(img_data)
                img = XlImage(img_buf)
                img.width = 600
                img.height = 350
                ws_charts.add_image(img, f"A{chart_row}")
                
                ws_charts.cell(row=chart_row + 20, column=1,
                              value=chart.get("title", "")).font = Font(
                    italic=True, size=9, color="6B7280", name="Calibri"
                )
                chart_row += 22
            except Exception as e:
                logger.warning(f"Erro ao embutir gráfico no XLSX: {e}")
    
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def generate_txt(markdown_text: str, title: str = "Relatório") -> io.BytesIO:
    """Gera versão plain text profissional do relatório."""
    width = 72
    lines = [
        "╔" + "═" * width + "╗",
        "║" + title.upper().center(width) + "║",
        "╚" + "═" * width + "╝",
        "",
        f"  Gerado por Joel — Agente de Análise de Documentos",
        f"  Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        "",
        "─" * (width + 2),
        "",
    ]
    
    for line in markdown_text.split("\n"):
        clean = line.replace("**", "").replace("*", "")
        stripped = clean.strip()
        
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].upper()
            lines.append("")
            lines.append("┌" + "─" * (len(text) + 4) + "┐")
            lines.append("│  " + text + "  │")
            lines.append("└" + "─" * (len(text) + 4) + "┘")
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            text = stripped[3:]
            lines.append("")
            lines.append("━" * min(len(text) + 4, width))
            lines.append(f"  {text}")
            lines.append("━" * min(len(text) + 4, width))
        elif stripped.startswith("### "):
            text = stripped[4:]
            lines.append("")
            lines.append(f"  ▸ {text}")
            lines.append("  " + "─" * min(len(text) + 2, width - 2))
        elif stripped == "---" or stripped == "***":
            lines.append("─" * (width + 2))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            lines.append(f"    •  {stripped[2:]}")
        elif re.match(r'^\d+\.\s+', stripped):
            lines.append(f"    {stripped}")
        elif stripped.startswith("> "):
            lines.append(f"    │ {stripped[2:]}")
        elif stripped.startswith("|"):
            lines.append(f"  {stripped}")
        else:
            if len(stripped) > width:
                words = stripped.split()
                current = "  "
                for word in words:
                    if len(current) + len(word) + 1 > width:
                        lines.append(current)
                        current = "  " + word
                    else:
                        current += " " + word if current.strip() else "  " + word
                if current.strip():
                    lines.append(current)
            else:
                lines.append(f"  {clean}" if stripped else "")
    
    lines.extend([
        "",
        "─" * (width + 2),
        "",
        "  Este relatório foi gerado automaticamente por Joel.",
        "  Recomenda-se validação independente antes de decisões críticas.",
        "",
        "╔" + "═" * width + "╗",
        "║" + "Joel — Agente de Análise de Documentos".center(width) + "║",
        "╚" + "═" * width + "╝",
    ])
    
    content = "\n".join(lines)
    buffer = io.BytesIO()
    buffer.write(content.encode("utf-8"))
    buffer.seek(0)
    return buffer
